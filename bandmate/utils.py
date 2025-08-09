import io, re, datetime
from typing import Dict, Any

from PIL import Image
import pytesseract

# ------------------ OCR ------------------

def ocr_images_to_text(files) -> str:
    """
    Read uploaded images (Werkzeug FileStorage list) and return combined OCR text.
    Adds simple [Page n] headers between pages.
    """
    pages = []
    for i, fs in enumerate(files, start=1):
        # Make sure the stream is at the start
        try:
            fs.stream.seek(0)
        except Exception:
            pass

        # Open image safely
        try:
            img = Image.open(fs.stream)
        except Exception:
            # Fallback: read bytes then open
            fs.stream.seek(0)
            data = fs.stream.read()
            img = Image.open(io.BytesIO(data))

        # Light preprocessing: grayscale only (fast & safe)
        img = img.convert("L")

        # OCR
        text = pytesseract.image_to_string(img, lang="eng").strip()
        pages.append(f"[Page {i}]\n{text}")

    out = "\n\n".join(pages).strip()
    if not out:
        raise RuntimeError("No text could be extracted from the uploaded image(s).")
    return out

# ------------------ Heuristics ------------------

def detect_task_type(text: str) -> str:
    t = text.lower()
    if any(k in t for k in ["diagram", "chart", "table", "map", "process", "figure"]):
        return "Task 1"
    return "Task 2"

def _band(x: float) -> float:
    # round to nearest .5 within 0–9
    try:
        return max(0.0, min(9.0, round(float(x) * 2) / 2))
    except Exception:
        return 6.0

# ------------------ Feedback (simple rule-based) ------------------

def generate_feedback(text: str, task_type: str, notes: str) -> Dict[str, Any]:
    words = len(text.split())
    bands = {"task": 6.0, "coherence": 6.0, "lexical": 6.0, "grammar": 6.0}
    actions = []

    if task_type == "Task 2" and words < 250:
        bands["task"] = 5.0
        actions.append("- Develop ideas further; aim for at least 250 words.")
    if re.search(r"\bi\b", text):  # lowercase 'i' pronoun
        bands["grammar"] = 5.5
        actions.append("- Capitalize the pronoun 'I'.")

    if not actions:
        actions = [
            "- Use clear topic sentences for each paragraph.",
            "- Add linking devices naturally (however, moreover, as a result).",
            "- Vary sentence structures (simple/compound/complex).",
            "- Check articles and subject–verb agreement.",
        ]

    if notes.strip():
        actions.append(f"- Teacher note: {notes.strip()}")

    summary = (
        f"{task_type} draft (~{words} words). Strengths and priorities across "
        f"idea development, coherence, vocabulary, and accuracy."
    )
    overall = _band(sum(bands.values()) / 4)
    return {"summary": summary, "actions": "\n".join(actions), "bands": {**bands, "overall": overall}}

# ------------------ DOCX / PDF builders ------------------

from docx import Document
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import cm

def build_docx(student_name: str, task_type: str, text: str, fb: Dict[str, Any]) -> bytes:
    doc = Document()
    doc.add_heading(f"BandMate — IELTS {task_type} Feedback", 1)
    if student_name:
        doc.add_paragraph(f"Student: {student_name}")
    doc.add_paragraph(datetime.datetime.now().strftime("%Y-%m-%d %H:%M"))

    doc.add_heading("Summary", 2)
    doc.add_paragraph(fb["summary"])

    b = fb["bands"]
    doc.add_heading("Band Estimates", 2)
    doc.add_paragraph(f"Task Response/Achievement: {b['task']}")
    doc.add_paragraph(f"Coherence & Cohesion: {b['coherence']}")
    doc.add_paragraph(f"Lexical Resource: {b['lexical']}")
    doc.add_paragraph(f"Grammatical Range & Accuracy: {b['grammar']}")
    doc.add_paragraph(f"Overall (weighted): {b['overall']}")

    doc.add_heading("Actionable Suggestions", 2)
    for line in fb["actions"].splitlines():
        if line.strip():
            p = doc.add_paragraph(line.strip()); p.style = "List Bullet"

    doc.add_heading("Student Text (OCR)", 2)
    doc.add_paragraph(text)

    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()

def build_pdf(student_name: str, task_type: str, text: str, fb: Dict[str, Any]) -> bytes:
    bio = io.BytesIO()
    c = canvas.Canvas(bio, pagesize=A4)
    w, h = A4
    x, y = 2 * cm, h - 2 * cm

    def ln(s, lead=14, bold=False):
        nonlocal y
        if y < 2 * cm:
            c.showPage(); y = h - 2 * cm
        c.setFont("Helvetica-Bold" if bold else "Helvetica", 12 if bold else 10)
        c.drawString(x, y, s); y -= lead

    ln(f"BandMate — IELTS {task_type} Feedback", 20, True)
    if student_name: ln(f"Student: {student_name}", 14)
    ln(datetime.datetime.now().strftime("%Y-%m-%d %H:%M"), 12)

    ln("Summary", 16, True)
    for t in fb["summary"].splitlines(): ln(t)

    b = fb["bands"]
    ln("Band Estimates", 16, True)
    for k, label in [
        ("task", "Task Response/Achievement"),
        ("coherence", "Coherence & Cohesion"),
        ("lexical", "Lexical Resource"),
        ("grammar", "Grammatical Range & Accuracy"),
        ("overall", "Overall (weighted)"),
    ]:
        ln(f"{label}: {b[k]}")

    ln("Actionable Suggestions", 16, True)
    for t in fb["actions"].splitlines(): ln(t)

    ln("Student Text (OCR)", 16, True)
    c.setFont("Helvetica", 10)
    for t in text.splitlines(): ln(t, 12)

    c.showPage(); c.save()
    return bio.getvalue()

def weighted_overall(b: dict, w_task: float, w_coh: float, w_lex: float, w_gra: float) -> float:
    tot = max(0.0001, w_task + w_coh + w_lex + w_gra)
    raw = (b["task"]*w_task + b["coherence"]*w_coh + b["lexical"]*w_lex + b["grammar"]*w_gra) / tot
    return _band(raw)
